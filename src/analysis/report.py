"""
Report generation for analysis results.
"""

import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import Wd_Table_Alignment
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportGenerator:
    """
    Generate analysis reports in various formats.
    """

    def __init__(self, config: Dict[str, Any]):
        """
        Initialize report generator.

        Args:
            config: Configuration dictionary
        """
        self.config = config

    def generate(
        self,
        results: Dict[str, Any],
        output_path: Union[str, Path],
        format: str = "docx",
    ) -> str:
        """
        Generate report from analysis results.

        Args:
            results: Dictionary of analysis results
            output_path: Output directory
            format: Report format ('docx', 'html', 'markdown')

        Returns:
            Path to generated report
        """
        output_path = Path(output_path)
        output_path.mkdir(parents=True, exist_ok=True)

        if format == "docx":
            return self._generate_docx(results, output_path)
        elif format == "html":
            return self._generate_html(results, output_path)
        elif format == "markdown":
            return self._generate_markdown(results, output_path)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_docx(
        self,
        results: Dict[str, Any],
        output_path: Path,
    ) -> str:
        """Generate Word document report."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word report generation")

        doc = Document()

        # Title
        doc.add_heading("Multi-Organ Segmentation Analysis Report", 0)

        # Patient info (if available)
        if "patient_info" in results:
            doc.add_heading("Patient Information", level=1)
            for key, value in results["patient_info"].items():
                doc.add_paragraph(f"{key}: {value}")

        # SUV Analysis
        if "suv" in results:
            doc.add_heading("SUV Analysis", level=1)
            self._add_suv_table(doc, results["suv"])

        # TMTV Analysis
        if "tmtv" in results:
            doc.add_heading("TMTV Analysis", level=1)
            self._add_tmtv_section(doc, results["tmtv"])

        # Histograms
        if "histogram" in results:
            doc.add_heading("Histogram Analysis", level=1)
            self._add_histogram_section(doc, output_path)

        # Save document
        report_path = output_path / "analysis_report.docx"
        doc.save(str(report_path))

        return str(report_path)

    def _add_suv_table(self, doc: Document, suv_results: Dict[str, Any]) -> None:
        """Add SUV analysis table to document."""
        if "organs" not in suv_results:
            return

        organs = suv_results["organs"]
        if not organs:
            return

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header
        header_cells = table.rows[0].cells
        headers = ["Organ", "Mean SUV", "Max SUV", "Volume (ml)", "SUV 50% Vol (ml)"]
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Data rows
        for organ in organs:
            row_cells = table.add_row().cells
            row_cells[0].text = organ.get("organ", "").replace("_", " ").title()
            row_cells[1].text = f"{organ.get('suv_mean', 0):.2f}"
            row_cells[2].text = f"{organ.get('suv_max', 0):.2f}"
            row_cells[3].text = f"{organ.get('volume_ml', 0):.2f}"
            row_cells[4].text = f"{organ.get('suv_50_volume', 0):.2f}"

    def _add_tmtv_section(self, doc: Document, tmtv_results: Dict[str, Any]) -> None:
        """Add TMTV analysis section to document."""
        for method, data in tmtv_results.items():
            if not isinstance(data, dict):
                continue

            doc.add_heading(f"Method: {method.replace('_', ' ').title()}", level=2)

            for key, value in data.items():
                if isinstance(value, float):
                    doc.add_paragraph(f"{key}: {value:.4f}")
                else:
                    doc.add_paragraph(f"{key}: {value}")

    def _add_histogram_section(self, doc: Document, output_path: Path) -> None:
        """Add histogram images to document."""
        image_files = [
            "organ_histograms.png",
            "combined_histogram.png",
            "threshold_curves.png",
            "cumulative_distribution.png",
        ]

        for img_name in image_files:
            img_path = output_path / img_name
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6))
                doc.add_paragraph()

    def _generate_html(
        self,
        results: Dict[str, Any],
        output_path: Path,
    ) -> str:
        """Generate HTML report."""
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Analysis Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                h1 { color: #333; }
                h2 { color: #666; }
                table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #4CAF50; color: white; }
                tr:nth-child(even) { background-color: #f2f2f2; }
                img { max-width: 100%; height: auto; margin: 20px 0; }
            </style>
        </head>
        <body>
            <h1>Multi-Organ Segmentation Analysis Report</h1>
        """

        # SUV Analysis
        if "suv" in results and "organs" in results["suv"]:
            html_content += "<h2>SUV Analysis</h2>"
            html_content += "<table><tr><th>Organ</th><th>Mean SUV</th><th>Max SUV</th><th>Volume (ml)</th></tr>"

            for organ in results["suv"]["organs"]:
                html_content += f"""
                <tr>
                    <td>{organ.get('organ', '').replace('_', ' ').title()}</td>
                    <td>{organ.get('suv_mean', 0):.2f}</td>
                    <td>{organ.get('suv_max', 0):.2f}</td>
                    <td>{organ.get('volume_ml', 0):.2f}</td>
                </tr>
                """

            html_content += "</table>"

        # TMTV Analysis
        if "tmtv" in results:
            html_content += "<h2>TMTV Analysis</h2>"
            for method, data in results["tmtv"].items():
                if isinstance(data, dict):
                    html_content += f"<h3>{method.replace('_', ' ').title()}</h3>"
                    html_content += "<ul>"
                    for key, value in data.items():
                        if isinstance(value, float):
                            html_content += f"<li>{key}: {value:.4f}</li>"
                        else:
                            html_content += f"<li>{key}: {value}</li>"
                    html_content += "</ul>"

        # Images
        image_files = ["organ_histograms.png", "combined_histogram.png", "threshold_curves.png"]
        for img_name in image_files:
            img_path = output_path / img_name
            if img_path.exists():
                html_content += f'<img src="{img_name}" alt="{img_name}">'

        html_content += "</body></html>"

        report_path = output_path / "analysis_report.html"
        with open(report_path, "w") as f:
            f.write(html_content)

        return str(report_path)

    def _generate_markdown(
        self,
        results: Dict[str, Any],
        output_path: Path,
    ) -> str:
        """Generate Markdown report."""
        md_content = "# Multi-Organ Segmentation Analysis Report\n\n"

        # SUV Analysis
        if "suv" in results and "organs" in results["suv"]:
            md_content += "## SUV Analysis\n\n"
            md_content += "| Organ | Mean SUV | Max SUV | Volume (ml) |\n"
            md_content += "|-------|----------|---------|-------------|\n"

            for organ in results["suv"]["organs"]:
                md_content += f"| {organ.get('organ', '').replace('_', ' ').title()} "
                md_content += f"| {organ.get('suv_mean', 0):.2f} "
                md_content += f"| {organ.get('suv_max', 0):.2f} "
                md_content += f"| {organ.get('volume_ml', 0):.2f} |\n"

            md_content += "\n"

        # TMTV Analysis
        if "tmtv" in results:
            md_content += "## TMTV Analysis\n\n"
            for method, data in results["tmtv"].items():
                if isinstance(data, dict):
                    md_content += f"### {method.replace('_', ' ').title()}\n\n"
                    for key, value in data.items():
                        if isinstance(value, float):
                            md_content += f"- **{key}**: {value:.4f}\n"
                        else:
                            md_content += f"- **{key}**: {value}\n"
                    md_content += "\n"

        # Images
        md_content += "## Visualizations\n\n"
        image_files = ["organ_histograms.png", "combined_histogram.png", "threshold_curves.png"]
        for img_name in image_files:
            img_path = output_path / img_name
            if img_path.exists():
                md_content += f"![{img_name}]({img_name})\n\n"

        report_path = output_path / "analysis_report.md"
        with open(report_path, "w") as f:
            f.write(md_content)

        return str(report_path)
